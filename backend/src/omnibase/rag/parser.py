"""Document parser: extract full text + structure from files.

Phase 1 upgrade from Phase 0 (which only extracted metadata).
Now extracts the actual text content with page/position tracking
for citation backlinks.

Supported formats:
- PDF: pypdf (text + page numbers)
- DOCX: python-docx (paragraphs + heading levels)
- TXT/MD: direct read (preserve line structure)

Output: ParsedDocument with pages, each containing text + page_number.
The chunker (B2) will further split these into embedding-sized chunks.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field

from omnibase.core.logging import get_logger

log = get_logger(__name__)


@dataclass
class ParsedPage:
    """A single page/section of a parsed document."""

    page_number: int
    text: str
    char_offset: int = 0  # offset from start of document (for citation backlinks)


@dataclass
class ParsedDocument:
    """A fully parsed document ready for chunking."""

    filename: str
    mime_type: str
    pages: list[ParsedPage] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        """Concatenated text of all pages."""
        return "\n\n".join(p.text for p in self.pages)

    @property
    def total_chars(self) -> int:
        return sum(len(p.text) for p in self.pages)


def parse_document(
    data: bytes,
    filename: str,
    mime_type: str,
) -> ParsedDocument:
    """Parse a document byte string into structured text.

    Routes to the appropriate parser based on mime type.
    Returns a ParsedDocument with at least one page.
    """
    if mime_type == "application/pdf":
        return _parse_pdf(data, filename)
    if mime_type == ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
        return _parse_docx(data, filename)
    if mime_type in ("text/plain", "text/markdown"):
        return _parse_text(data, filename, mime_type)
    # Fallback: try as text
    log.warning("parser.unknown_mime_type", mime_type=mime_type, filename=filename)
    return _parse_text(data, filename, mime_type)


def _parse_pdf(data: bytes, filename: str) -> ParsedDocument:
    """Parse PDF using pypdf — extract text per page."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data), strict=False)
        pages: list[ParsedPage] = []
        char_offset = 0

        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception as exc:
                log.debug("parser.pdf.page_failed", page=i, error=str(exc))
                text = ""

            # Clean: collapse excessive whitespace but preserve paragraph breaks
            text = _clean_text(text)
            if text.strip():
                pages.append(
                    ParsedPage(
                        page_number=i + 1,
                        text=text,
                        char_offset=char_offset,
                    )
                )
                char_offset += len(text) + 2  # +2 for the \n\n separator

        # Extract metadata
        meta: dict = {}
        try:
            info = reader.metadata
            if info:
                for key in ["/Title", "/Author", "/Subject", "/Creator"]:
                    val = info.get(key)
                    if val:
                        meta[key.lstrip("/").lower()] = str(val).strip()
        except Exception:
            pass
        meta["page_count"] = len(reader.pages)

        log.info(
            "parser.pdf.complete",
            filename=filename,
            pages=len(pages),
            chars=sum(len(p.text) for p in pages),
        )
        return ParsedDocument(
            filename=filename,
            mime_type="application/pdf",
            pages=pages,
            metadata=meta,
        )
    except Exception as exc:
        log.error("parser.pdf.failed", filename=filename, error=str(exc))
        return ParsedDocument(
            filename=filename,
            mime_type="application/pdf",
            pages=[ParsedPage(page_number=1, text="")],
            metadata={"error": str(exc)[:500]},
        )


def _parse_docx(data: bytes, filename: str) -> ParsedDocument:
    """Parse DOCX using python-docx — extract paragraphs with heading levels."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Prefix headings for structure awareness
            paragraph_style = para.style
            style = (paragraph_style.name or "").lower() if paragraph_style is not None else ""
            if "heading 1" in style:
                paragraphs.append(f"\n# {text}\n")
            elif "heading 2" in style:
                paragraphs.append(f"\n## {text}\n")
            elif "heading 3" in style:
                paragraphs.append(f"\n### {text}\n")
            else:
                paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)
        meta = {"paragraph_count": len(paragraphs)}

        log.info(
            "parser.docx.complete",
            filename=filename,
            paragraphs=len(paragraphs),
            chars=len(full_text),
        )
        return ParsedDocument(
            filename=filename,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            pages=[ParsedPage(page_number=1, text=full_text, char_offset=0)],
            metadata=meta,
        )
    except ImportError:
        log.warning("parser.docx.not_installed", filename=filename)
        return _parse_text(data, filename, "text/plain")
    except Exception as exc:
        log.error("parser.docx.failed", filename=filename, error=str(exc))
        return ParsedDocument(
            filename=filename,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            pages=[ParsedPage(page_number=1, text="")],
            metadata={"error": str(exc)[:500]},
        )


def _parse_text(data: bytes, filename: str, mime_type: str) -> ParsedDocument:
    """Parse plain text or Markdown — direct decode."""
    try:
        # Try UTF-8 first, fallback to latin-1
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1", errors="replace")

        text = _clean_text(text)
        is_markdown = mime_type == "text/markdown" or filename.lower().endswith(".md")

        log.info(
            "parser.text.complete",
            filename=filename,
            chars=len(text),
            markdown=is_markdown,
        )
        return ParsedDocument(
            filename=filename,
            mime_type=mime_type,
            pages=[ParsedPage(page_number=1, text=text, char_offset=0)],
            metadata={"markdown": is_markdown, "char_count": len(text)},
        )
    except Exception as exc:
        log.error("parser.text.failed", filename=filename, error=str(exc))
        return ParsedDocument(
            filename=filename,
            mime_type=mime_type,
            pages=[ParsedPage(page_number=1, text="")],
            metadata={"error": str(exc)[:500]},
        )


def _clean_text(text: str) -> str:
    """Clean extracted text: collapse excessive whitespace, fix encoding artifacts."""
    import re

    # Collapse multiple blank lines to max 2
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    # Collapse multiple spaces (but not newlines)
    text = re.sub(r"[ \t]{3,}", "  ", text)
    # Remove common PDF extraction artifacts
    text = text.replace("\x00", "").replace("\ufffd", "")
    return text.strip()


__all__ = ["ParsedDocument", "ParsedPage", "parse_document"]
